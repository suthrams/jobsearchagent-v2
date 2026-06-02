"""Resume text/file renderer for the Resume Clinic (ADR-066 nice-to-have).

Renders a clinic review's final resume in six formats: Markdown, plain text,
HTML, JSON Resume, DOCX, and PDF. One canonical composer turns
`(parsed_profile, overhaul, edited, decision)` into a small intermediate
(`RenderedResume`); each render_* function walks that intermediate.

The renderer is deterministic — no LLM call. Fidelity is preserved by the
composer:

  - `decision == "reject"`   -> render the original resume unchanged.
  - `decision == "edit"`     -> apply `edited` (the human-authored overhaul).
                                ADR-059: the human is the final author and is
                                NOT re-policed by Fidelity.
  - `decision == "approve"`  -> apply the agent's `overhaul`.
  - `decision in (None, "revise")` -> preview the agent's `overhaul` and set
                                `RenderedResume.banner` so the renderers can
                                surface "this is a preview" on the output.

A rewrite that does not match any bullet in the source is appended at the end
of its section so nothing is silently dropped. Placeholder tokens emitted by
the agent (e.g. `[N]+ services`, `[X]%`) are kept verbatim — the candidate
fills them in.
"""

from __future__ import annotations

import html
import io
import json
import re
from dataclasses import dataclass, field
from typing import Iterable, Literal

# ── Intermediate representation ──────────────────────────────────────────────


@dataclass
class ExperienceItem:
    title: str
    company: str
    dates: str
    technologies: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)


@dataclass
class EducationItem:
    degree: str
    institution: str
    year: str | None = None
    gpa: str | None = None                   # ADR-067
    honors: list[str] = field(default_factory=list)  # ADR-067


@dataclass
class CertificationItem:
    name: str
    issuer: str | None = None
    year: str | None = None


@dataclass
class SkillGroupItem:
    category: str
    skills: list[str] = field(default_factory=list)


@dataclass
class RenderedResume:
    name: str | None = None
    headline: str | None = None
    email: str | None = None
    location: str | None = None
    summary: str | None = None
    experience: list[ExperienceItem] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    # ADR-067: when populated, the renderer uses skill_groups for the Skills
    # section instead of the flat list. Both views co-exist so older
    # parsed_profile rows that lack groupings still render via `skills`.
    skill_groups: list[SkillGroupItem] = field(default_factory=list)
    education: list[EducationItem] = field(default_factory=list)
    certifications: list[CertificationItem] = field(default_factory=list)
    # The order in which to render the four content sections. Each value must be
    # one of: "summary", "experience", "skills", "education", "certifications".
    section_order: list[str] = field(default_factory=lambda: [
        "summary", "experience", "skills", "certifications", "education",
    ])
    banner: str | None = None    # surfaces a "preview" notice when applicable


# ── Composer ─────────────────────────────────────────────────────────────────


def compose_resume(profile: dict,
                   overhaul: dict | None,
                   edited: dict | None,
                   decision: str | None) -> RenderedResume:
    """Apply the clinic overhaul (or its edited override) to the parsed profile
    per the decision; return the rendered intermediate.

    `profile` is the parsed `ResumeProfile` shape (dict, as persisted in
    `resumes.parsed_profile_json`). `overhaul` is the agent's
    `{reorganization, rewrites}`; `edited` is the optional human-authored
    overhaul of the same shape (used when decision == "edit").
    """
    base = _profile_to_rendered(profile)

    # 1) Pick which overhaul to apply, if any.
    #
    # ADR-068 rule: prefer `edited` whenever it is populated (the chat-revise
    # loop writes here per turn, regardless of the user's final decision),
    # EXCEPT when `decision == "reject"` - reject is the explicit "throw out
    # the overhaul" signal and falls back to the original parsed resume.
    applied: dict | None
    banner: str | None = None
    d = (decision or "").strip().lower() or None

    if d == "reject":
        applied = None
    elif edited:
        applied = edited
        if d == "edit":
            banner = None
        elif d == "revise":
            banner = "Preview - editing in progress (decision: revise)."
        elif d == "approve":
            banner = "Preview - approved, with chat edits applied."
        elif d is None:
            banner = "Preview - editing in progress (no decision yet)."
        else:
            banner = f"Preview - editing in progress (decision: {decision!r})."
    elif d == "approve":
        applied = overhaul
    elif d == "revise":
        applied = overhaul
        banner = "Preview - decision is 'revise'."
    elif d is None:
        applied = overhaul
        if applied:
            banner = "Preview - no decision recorded yet."
    else:
        # Unknown decision values are treated as preview rather than dropped.
        applied = overhaul
        if applied:
            banner = f"Preview - unrecognized decision: {decision!r}."

    if applied:
        _apply_rewrites(base, applied.get("rewrites") or [])
        _apply_reorganization(base, applied.get("reorganization") or {})

    base.banner = banner
    return base


def _profile_to_rendered(profile: dict) -> RenderedResume:
    """Materialise a ResumeProfile dict into the RenderedResume shape."""
    experience: list[ExperienceItem] = []
    for entry in profile.get("experience") or []:
        if not isinstance(entry, dict):
            continue
        dates = _format_dates(entry.get("start_year"), entry.get("end_year"))
        experience.append(ExperienceItem(
            title=str(entry.get("title") or "").strip(),
            company=str(entry.get("company") or "").strip(),
            dates=dates,
            technologies=[str(t) for t in (entry.get("technologies") or [])],
            bullets=_description_to_bullets(entry.get("description")),
        ))

    education: list[EducationItem] = []
    for e in profile.get("education") or []:
        if not isinstance(e, dict):
            continue
        # ADR-067: also pull gpa + honors when present.
        gpa = e.get("gpa")
        gpa_str = str(gpa).strip() if gpa not in (None, "") else None
        honors_raw = e.get("honors") or []
        honors = [str(h).strip() for h in honors_raw if str(h).strip()] if isinstance(honors_raw, list) else []
        education.append(EducationItem(
            degree=str(e.get("degree") or "").strip(),
            institution=str(e.get("institution") or "").strip(),
            year=_year_str(e.get("year")),
            gpa=gpa_str,
            honors=honors,
        ))

    certifications: list[CertificationItem] = []
    for c in profile.get("certifications") or []:
        if not isinstance(c, dict):
            continue
        certifications.append(CertificationItem(
            name=str(c.get("name") or "").strip(),
            issuer=str(c.get("issuer") or "").strip() or None,
            year=_year_str(c.get("year")),
        ))

    # ADR-067: skill_groups overlays the flat skills list. Render-time logic
    # in the format-specific renderers prefers `skill_groups` when populated.
    skill_groups: list[SkillGroupItem] = []
    for g in profile.get("skill_groups") or []:
        if not isinstance(g, dict):
            continue
        cat = str(g.get("category") or "").strip()
        items = [str(s).strip() for s in (g.get("skills") or []) if str(s).strip()]
        if cat and items:
            skill_groups.append(SkillGroupItem(category=cat, skills=items))

    return RenderedResume(
        name=(profile.get("name") or "").strip() or None,
        headline=(profile.get("headline") or "").strip() or None,
        email=(profile.get("email") or "").strip() or None,
        location=(profile.get("location") or "").strip() or None,
        summary=(profile.get("summary") or "").strip() or None,
        experience=experience,
        skills=[str(s).strip() for s in (profile.get("skills") or []) if str(s).strip()],
        skill_groups=skill_groups,
        education=education,
        certifications=certifications,
    )


def _format_dates(start: object, end: object) -> str:
    s = _year_str(start) or ""
    e = _year_str(end)
    if not s and not e:
        return ""
    if not e:
        return f"{s} - present" if s else ""
    if not s:
        return str(e)
    return f"{s} - {e}"


def _year_str(value: object) -> str | None:
    if value is None or value == "":
        return None
    try:
        return str(int(value))
    except (TypeError, ValueError):
        return str(value).strip() or None


_BULLET_MARKERS = ("•", "·", "*", "-", "–", "—", "►", "▪", "◦")
_MIN_SENTENCE_LEN = 25       # below this, treat as a fragment (don't promote)
_PARA_SPLIT_THRESHOLD = 100  # only sentence-split paragraphs above this length


def _description_to_bullets(description: object) -> list[str]:
    """Split an experience description into bullets.

    Resume parses are noisy. The default behaviour matters because it is the
    single biggest driver of output quality. The contract, in order of
    precedence:

    1. Explicit bullet markers (>=2 lines start with a `•`/`-`/`*` etc):
       parse as a bullet list; continuation lines fold into the bullet above.
    2. Multiple non-empty lines without bullet markers: each line is a bullet
       (assumes the parser already split the bullets with newlines).
    3. Single paragraph >100 chars AND looks like multiple full sentences:
       sentence-split on a clean regex. Fragments below 25 chars are folded
       back into the previous bullet so half-sentences never stand alone.
    4. Otherwise: one bullet, intact.

    The trade-off in case 3 is fidelity vs readability. Rendering a 400-char
    paragraph as one giant bullet is unreadable; sentence-splitting produces
    2-4 bullets of complete sentences, which is what readers expect on a
    resume.
    """
    if not description:
        return []
    text = str(description).strip()
    if not text:
        return []

    lines = text.splitlines()
    nonempty_lines = [l for l in lines if l.strip()]

    # Case 1: explicit bullet markers (>=2 bulleted lines).
    bulleted = [l for l in nonempty_lines
                if l.lstrip()[0:1] in _BULLET_MARKERS]
    if len(bulleted) >= 2:
        out: list[str] = []
        current: list[str] = []
        for line in nonempty_lines:
            stripped = line.lstrip()
            if stripped[0:1] in _BULLET_MARKERS:
                if current:
                    out.append(" ".join(current).strip())
                rest = re.sub(r"^[\-\*•·–—►▪◦]+\s*", "", stripped)
                current = [rest] if rest else []
            else:
                if current:
                    current.append(stripped.strip())
                else:
                    out.append(stripped.strip())
        if current:
            out.append(" ".join(current).strip())
        return [b for b in out if b]

    # Case 2: multiple non-empty lines without markers.
    if len(nonempty_lines) > 1:
        return [l.strip() for l in nonempty_lines]

    # Case 3: single paragraph above the split threshold with multiple
    # sentences -> sentence-split. The regex splits after `.`/`!`/`?` followed
    # by whitespace AND a capital letter / digit / opening bracket, so common
    # abbreviations (e.g. "U.S.") don't trigger a split.
    if len(text) >= _PARA_SPLIT_THRESHOLD and text.count(".") >= 1:
        parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\[])", text)
        sentences = [p.strip() for p in parts if p.strip()]
        if len(sentences) >= 2:
            # Re-merge fragments shorter than the floor into the prior bullet.
            merged: list[str] = []
            for s in sentences:
                if merged and len(s) < _MIN_SENTENCE_LEN:
                    merged[-1] = f"{merged[-1]} {s}"
                else:
                    merged.append(s)
            return merged

    # Case 4: keep as one bullet.
    return [text]


# ── Rewrite + reorganization application ────────────────────────────────────


# Maps "experience:Company:Title" → (company, title). Whitespace-tolerant; case-
# insensitive matching used inside _find_experience_item.
def _parse_experience_label(label: str) -> tuple[str, str] | None:
    if not isinstance(label, str):
        return None
    parts = label.split(":", 2)
    if len(parts) < 3 or parts[0].strip().lower() != "experience":
        return None
    return parts[1].strip(), parts[2].strip()


def _find_experience_item(items: list[ExperienceItem],
                          company: str, title: str) -> ExperienceItem | None:
    co = company.casefold()
    ti = title.casefold()
    for it in items:
        if it.company.casefold() == co and it.title.casefold() == ti:
            return it
    return None


def _apply_rewrites(base: RenderedResume, rewrites: Iterable[dict]) -> None:
    """For each rewrite, find the bullet it points at (by section_label +
    original_text) and substitute. If no match, append at the end of the
    section so the rewrite is never silently dropped.
    """
    for r in rewrites or []:
        if not isinstance(r, dict):
            continue
        label = (r.get("section_label") or "").strip()
        suggested = (r.get("suggested_text") or "").strip()
        original = (r.get("original_text") or "").strip()
        if not suggested:
            continue
        if label == "headline":
            _replace_or_append_headline(base, original, suggested)
            continue
        if label == "summary":
            _replace_or_append_summary(base, original, suggested)
            continue
        if label == "skills":
            # A clinic rewrite at section_label="skills" is rare; treat the
            # suggested text as a single skill addition.
            if suggested and suggested not in base.skills:
                base.skills.append(suggested)
            continue
        parsed = _parse_experience_label(label)
        if parsed is not None:
            company, title = parsed
            item = _find_experience_item(base.experience, company, title)
            if item is not None:
                _replace_or_append_bullet(item.bullets, original, suggested)
                continue
        # Unknown / unmatched label: keep the rewrite as an appended bullet on
        # the most-recent experience entry rather than dropping it.
        if base.experience:
            base.experience[0].bullets.append(suggested)


def _replace_or_append_headline(base: RenderedResume, original: str, suggested: str) -> None:
    """Apply a `section_label="headline"` rewrite. The headline is a single
    line (not bullets), so the substitution is whole-string with the same
    exact-then-substring fallback summary uses."""
    if base.headline and original and original in base.headline:
        base.headline = base.headline.replace(original, suggested, 1)
    else:
        base.headline = suggested


def _replace_or_append_summary(base: RenderedResume, original: str, suggested: str) -> None:
    if base.summary and original and original in base.summary:
        base.summary = base.summary.replace(original, suggested, 1)
    elif base.summary and not original:
        base.summary = f"{base.summary}\n\n{suggested}"
    else:
        base.summary = suggested


def _replace_or_append_bullet(bullets: list[str], original: str, suggested: str) -> None:
    if not original:
        bullets.append(suggested)
        return
    # Exact match first.
    for i, b in enumerate(bullets):
        if b.strip() == original:
            bullets[i] = suggested
            return
    # Substring fallback - many real rewrites pull a clean version of a noisy bullet.
    o = original.casefold()
    for i, b in enumerate(bullets):
        if o and o in b.casefold():
            bullets[i] = suggested
            return
    bullets.append(suggested)


def _apply_reorganization(base: RenderedResume, reorg: dict) -> None:
    """`section_order` is a list of section names. `moves` is a list of
    {action: move|cut|promote, subject: str, rationale: str}.

    For v1 we honor:
      - the proposed `section_order` (filtered to known sections)
      - `cut` moves whose subject names a section (e.g. "Education")
      - `promote` moves whose subject names a section (it floats to the top of
        the body, just after summary)

    Per-bullet moves (e.g. moving an experience bullet between roles) are not
    in v1 because the section_label vocabulary doesn't carry an action for
    them; the agent already emits those via `rewrites`.
    """
    raw_order = reorg.get("section_order") or []
    valid = {"summary", "experience", "skills", "education", "certifications"}
    proposed: list[str] = []
    for s in raw_order:
        norm = str(s).strip().lower()
        # Allow tokens like "header / contact" or "projects" to pass through
        # the order list - they map to nothing in v1, but their presence in
        # the iteration order doesn't break anything; we just skip them when
        # rendering. The fix is: only KEEP tokens we know how to render.
        if norm in valid and norm not in proposed:
            proposed.append(norm)
    if proposed:
        # Agents commonly omit "summary" from their proposed section_order on
        # the assumption that it sits implicitly above the body, in the
        # header. The renderer treats summary as a body section, so an omitted
        # summary needs to be RE-INSERTED above the agent's proposed order
        # (not appended at the end - that was the original bug).
        if "summary" not in proposed:
            proposed.insert(0, "summary")
        # Append any remaining known sections at the END so they aren't
        # silently dropped.
        for s in ("summary", "experience", "skills", "certifications", "education"):
            if s not in proposed:
                proposed.append(s)
        base.section_order = proposed

    moves = reorg.get("moves") or []
    for m in moves:
        if not isinstance(m, dict):
            continue
        action = (m.get("action") or "").strip().lower()
        subject = (m.get("subject") or "").strip()
        if not subject:
            continue
        norm_subject = subject.casefold()

        if action == "cut":
            if norm_subject in {"summary", "experience", "skills",
                                "education", "certifications"}:
                # Drop that whole section by removing it from section_order.
                base.section_order = [s for s in base.section_order
                                      if s != norm_subject]
            else:
                # Try to cut a certification by its name.
                base.certifications = [c for c in base.certifications
                                       if c.name.casefold() != norm_subject]

        elif action == "promote":
            if norm_subject in {"summary", "experience", "skills",
                                "education", "certifications"}:
                base.section_order = [s for s in base.section_order
                                      if s != norm_subject]
                # Insert just after summary (or at front if summary not shown).
                insert_at = 1 if base.section_order and base.section_order[0] == "summary" else 0
                base.section_order.insert(insert_at, norm_subject)
        # `move` is a no-op in v1; the agent's rewrites already represent
        # bullet-level moves.


# ── Renderers ────────────────────────────────────────────────────────────────


# ── Markdown (canonical) ────────────────────────────────────────────────────

def render_markdown(rendered: RenderedResume) -> str:
    lines: list[str] = []
    if rendered.name:
        lines.append(f"# {rendered.name}")
    if rendered.headline:
        lines.append(f"_{rendered.headline}_")
    contact_bits = [b for b in (rendered.email, rendered.location) if b]
    if contact_bits:
        lines.append(" · ".join(contact_bits))
    if rendered.name or rendered.headline or contact_bits:
        lines.append("")

    for section in rendered.section_order:
        if section == "summary" and rendered.summary:
            lines.append("## Summary")
            lines.append("")
            for para in rendered.summary.split("\n\n"):
                lines.append(para.strip())
                lines.append("")
        elif section == "experience" and rendered.experience:
            lines.append("## Experience")
            lines.append("")
            for it in rendered.experience:
                # `Title · Company    dates` on the heading line.
                title_co = " · ".join(b for b in (it.title, it.company) if b)
                if title_co and it.dates:
                    lines.append(f"### {title_co} — {it.dates}")
                elif title_co:
                    lines.append(f"### {title_co}")
                elif it.dates:
                    lines.append(f"### {it.dates}")
                if it.technologies:
                    lines.append(f"*{', '.join(it.technologies)}*")
                lines.append("")
                for b in it.bullets:
                    lines.append(f"- {b}")
                lines.append("")
        elif section == "skills" and (rendered.skill_groups or rendered.skills):
            lines.append("## Skills")
            # ADR-067: render grouped when categorised data is present;
            # otherwise fall back to the flat list.
            if rendered.skill_groups:
                for g in rendered.skill_groups:
                    lines.append(f"**{g.category}:** {' · '.join(g.skills)}")
            else:
                lines.append(" · ".join(rendered.skills))
            lines.append("")
        elif section == "certifications" and rendered.certifications:
            lines.append("## Certifications")
            for c in rendered.certifications:
                meta = ", ".join(b for b in (c.issuer, c.year) if b)
                lines.append(f"- {c.name}" + (f" ({meta})" if meta else ""))
            lines.append("")
        elif section == "education" and rendered.education:
            lines.append("## Education")
            for ed in rendered.education:
                # ADR-067: include GPA and honors when the parser captured them.
                head_bits = [b for b in (ed.degree, ed.institution, ed.year) if b]
                head_line = f"- {' · '.join(head_bits)}"
                if ed.gpa:
                    head_line += f" — GPA: {ed.gpa}"
                lines.append(head_line)
                for h in ed.honors:
                    lines.append(f"  - {h}")
            lines.append("")

    # Discreet "preview" footer instead of a loud banner up top.
    if rendered.banner:
        lines.append("")
        lines.append("---")
        lines.append(f"_{rendered.banner}_")

    # Collapse runs of blank lines and trailing whitespace.
    out: list[str] = []
    prev_blank = False
    for ln in lines:
        is_blank = (ln.strip() == "")
        if is_blank and prev_blank:
            continue
        out.append(ln.rstrip())
        prev_blank = is_blank
    return "\n".join(out).strip() + "\n"


# ── Plain text (markdown stripped + ATS-friendly headers) ────────────────────

def render_plain_text(rendered: RenderedResume, *, width: int = 72) -> str:
    lines: list[str] = []
    sep = "-" * 60
    if rendered.name:
        lines.append(rendered.name.upper())
    if rendered.headline:
        lines.extend(_wrap(rendered.headline, width=width))
    contact_bits = [b for b in (rendered.email, rendered.location) if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    if rendered.name or rendered.headline or contact_bits:
        lines.append("")

    def _section_header(title: str) -> None:
        lines.append(sep)
        lines.append(title.upper())
        lines.append(sep)

    for section in rendered.section_order:
        if section == "summary" and rendered.summary:
            _section_header("Summary")
            for para in rendered.summary.split("\n\n"):
                lines.extend(_wrap(para.strip(), width=width))
                lines.append("")
        elif section == "experience" and rendered.experience:
            _section_header("Experience")
            for it in rendered.experience:
                head = " | ".join(b for b in (it.title, it.company) if b)
                # Right-align the dates on the same line if there is room.
                if head and it.dates and len(head) + len(it.dates) + 2 <= width:
                    pad = " " * (width - len(head) - len(it.dates))
                    lines.append(f"{head}{pad}{it.dates}")
                else:
                    if head:
                        lines.append(head)
                    if it.dates:
                        lines.append(it.dates)
                if it.technologies:
                    lines.append(f"  Technologies: {', '.join(it.technologies)}")
                for b in it.bullets:
                    bullet_lines = _wrap(b, width=width - 4)
                    if bullet_lines:
                        lines.append(f"  * {bullet_lines[0]}")
                        for sub in bullet_lines[1:]:
                            lines.append(f"    {sub}")
                lines.append("")
        elif section == "skills" and (rendered.skill_groups or rendered.skills):
            _section_header("Skills")
            # ADR-067: grouped when available.
            if rendered.skill_groups:
                for g in rendered.skill_groups:
                    line = f"{g.category}: {', '.join(g.skills)}"
                    lines.extend(_wrap(line, width=width))
            else:
                lines.extend(_wrap(", ".join(rendered.skills), width=width))
            lines.append("")
        elif section == "certifications" and rendered.certifications:
            _section_header("Certifications")
            for c in rendered.certifications:
                meta = ", ".join(b for b in (c.issuer, c.year) if b)
                lines.append(f"{c.name}" + (f", {meta}" if meta else ""))
            lines.append("")
        elif section == "education" and rendered.education:
            _section_header("Education")
            for ed in rendered.education:
                head_bits = [b for b in (ed.degree, ed.institution, ed.year) if b]
                head = ", ".join(head_bits)
                if ed.gpa:
                    head = f"{head} | GPA: {ed.gpa}"
                lines.append(head)
                for h in ed.honors:
                    lines.append(f"  * {h}")
            lines.append("")

    # Discreet preview footer if applicable.
    if rendered.banner:
        if lines and lines[-1].strip() != "":
            lines.append("")
        lines.append(f"({rendered.banner})")
    # Collapse trailing blanks.
    while lines and lines[-1].strip() == "":
        lines.pop()
    return "\n".join(lines) + "\n"


def _wrap(text: str, *, width: int) -> list[str]:
    text = text.strip()
    if not text:
        return []
    words = text.split()
    out: list[str] = []
    line = ""
    for w in words:
        if not line:
            line = w
        elif len(line) + 1 + len(w) <= width:
            line = f"{line} {w}"
        else:
            out.append(line)
            line = w
    if line:
        out.append(line)
    return out


# ── HTML (hand-rolled; no extra dependency) ──────────────────────────────────

def render_html(rendered: RenderedResume) -> str:
    e = html.escape
    parts: list[str] = []
    parts.append("<!DOCTYPE html>")
    parts.append('<html lang="en"><head><meta charset="utf-8">')
    parts.append(f"<title>{e(rendered.name or 'Resume')}</title>")
    parts.append("<style>")
    parts.append(
        "body{font-family:'Helvetica Neue',Helvetica,Arial,sans-serif;"
        "max-width:780px;margin:36px auto;padding:0 28px;color:#222;"
        "line-height:1.4;font-size:13.5px}"
        "h1{margin:0 0 2px 0;font-size:26px;letter-spacing:.2px}"
        ".headline{color:#555;font-style:italic;margin-bottom:2px}"
        ".contact{color:#666;font-size:12.5px;margin:0 0 4px}"
        ".rule{border:0;border-top:1px solid #ddd;margin:10px 0 4px}"
        "h2{font-size:12.5px;text-transform:uppercase;letter-spacing:1.2px;"
        "color:#333;border-bottom:1px solid #ddd;padding-bottom:2px;"
        "margin:14px 0 6px}"
        ".role{display:flex;justify-content:space-between;"
        "align-items:baseline;margin:8px 0 0}"
        ".role .title{font-weight:600;font-size:13.5px}"
        ".role .dates{color:#666;font-size:12px;white-space:nowrap;margin-left:12px}"
        ".tech{color:#666;font-style:italic;font-size:12px;margin:0 0 4px}"
        "ul{margin:4px 0 8px 20px;padding:0}"
        "li{margin:2px 0}"
        ".skills{margin-bottom:6px}"
        ".preview{color:#888;font-style:italic;font-size:11.5px;"
        "border-top:1px dotted #ccc;margin-top:18px;padding-top:6px}"
    )
    parts.append("</style></head><body>")
    if rendered.name:
        parts.append(f"<h1>{e(rendered.name)}</h1>")
    if rendered.headline:
        parts.append(f'<div class="headline">{e(rendered.headline)}</div>')
    contact = " &middot; ".join(e(b) for b in (rendered.email, rendered.location) if b)
    if contact:
        parts.append(f'<div class="contact">{contact}</div>')
    if rendered.name or rendered.headline or contact:
        parts.append('<hr class="rule">')

    for section in rendered.section_order:
        if section == "summary" and rendered.summary:
            parts.append("<h2>Summary</h2>")
            for para in rendered.summary.split("\n\n"):
                parts.append(f"<p>{e(para.strip())}</p>")
        elif section == "experience" and rendered.experience:
            parts.append("<h2>Experience</h2>")
            for it in rendered.experience:
                head = " &middot; ".join(e(b) for b in (it.title, it.company) if b)
                if head or it.dates:
                    parts.append('<div class="role">')
                    if head:
                        parts.append(f'<span class="title">{head}</span>')
                    if it.dates:
                        parts.append(f'<span class="dates">{e(it.dates)}</span>')
                    parts.append("</div>")
                if it.technologies:
                    parts.append(f'<div class="tech">{e(", ".join(it.technologies))}</div>')
                if it.bullets:
                    parts.append("<ul>")
                    for b in it.bullets:
                        parts.append(f"<li>{e(b)}</li>")
                    parts.append("</ul>")
        elif section == "skills" and (rendered.skill_groups or rendered.skills):
            parts.append("<h2>Skills</h2>")
            if rendered.skill_groups:
                # ADR-067: one row per category.
                parts.append('<div class="skills">')
                for g in rendered.skill_groups:
                    parts.append(
                        f'<div><strong>{e(g.category)}:</strong> '
                        f'{" &middot; ".join(e(s) for s in g.skills)}</div>'
                    )
                parts.append("</div>")
            else:
                parts.append(
                    f'<div class="skills">'
                    f'{" &middot; ".join(e(s) for s in rendered.skills)}</div>'
                )
        elif section == "certifications" and rendered.certifications:
            parts.append("<h2>Certifications</h2>")
            parts.append("<ul>")
            for c in rendered.certifications:
                meta = ", ".join(b for b in (c.issuer, c.year) if b)
                txt = e(c.name) + (f' <span class="meta">({e(meta)})</span>' if meta else "")
                parts.append(f"<li>{txt}</li>")
            parts.append("</ul>")
        elif section == "education" and rendered.education:
            parts.append("<h2>Education</h2>")
            parts.append("<ul>")
            for ed in rendered.education:
                head_bits = [b for b in (ed.degree, ed.institution, ed.year) if b]
                head = e(" &middot; ".join(head_bits))
                if ed.gpa:
                    head += f' <span class="meta">&mdash; GPA: {e(ed.gpa)}</span>'
                if ed.honors:
                    inner_items = "".join(
                        f"<li>{e(h)}</li>" for h in ed.honors
                    )
                    parts.append(f"<li>{head}<ul>{inner_items}</ul></li>")
                else:
                    parts.append(f"<li>{head}</li>")
            parts.append("</ul>")

    if rendered.banner:
        parts.append(f'<div class="preview">{e(rendered.banner)}</div>')
    parts.append("</body></html>")
    return "\n".join(parts)


# ── JSON Resume (jsonresume.org schema subset) ───────────────────────────────

def render_json_resume(rendered: RenderedResume) -> dict:
    out: dict = {
        "$schema": "https://jsonresume.org/schema/1.0.0/resume.json",
        "basics": {
            "name": rendered.name or "",
            "label": rendered.headline or "",
            "email": rendered.email or "",
            "location": {"city": rendered.location or "", "region": ""},
            "summary": rendered.summary or "",
        },
        "work": [
            {
                "name": it.company,
                "position": it.title,
                "startDate": _start_from_dates(it.dates),
                "endDate": _end_from_dates(it.dates),
                "summary": "\n".join(it.bullets),
                "highlights": list(it.bullets),
                "keywords": list(it.technologies),
            }
            for it in rendered.experience
        ],
        # ADR-067: emit skill_groups under JSON Resume's optional category
        # via the `keywords` field of each skill entry. JSON Resume's schema
        # supports skill.name + skill.keywords; using `keywords` for the
        # category lets us preserve the grouping while staying within the
        # schema. The flat skills list is still each skill's `name`.
        "skills": (
            [{"name": s, "keywords": [g.category]}
             for g in rendered.skill_groups for s in g.skills]
            if rendered.skill_groups else
            [{"name": s} for s in rendered.skills]
        ),
        "education": [
            {
                "institution": ed.institution,
                "area": ed.degree,
                "endDate": ed.year or "",
                "score": ed.gpa or "",          # ADR-067
                "courses": list(ed.honors),     # ADR-067: honors stuffed in courses (schema-allowed list)
            }
            for ed in rendered.education
        ],
        "certificates": [
            {
                "name": c.name,
                "issuer": c.issuer or "",
                "date": c.year or "",
            }
            for c in rendered.certifications
        ],
        "meta": {
            "section_order": list(rendered.section_order),
            "banner": rendered.banner or "",
        },
    }
    return out


def _start_from_dates(dates: str) -> str:
    if not dates:
        return ""
    return dates.split(" - ", 1)[0].strip()


def _end_from_dates(dates: str) -> str:
    if not dates:
        return ""
    if " - " not in dates:
        return ""
    end = dates.split(" - ", 1)[1].strip()
    return "" if end.lower() == "present" else end


# ── DOCX (python-docx) ───────────────────────────────────────────────────────

def render_docx(rendered: RenderedResume) -> bytes:
    # Imported lazily so the module can be imported in environments that have
    # not installed python-docx (e.g. partial test runs).
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()
    # Tighter margins for a one-page resume.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)
    base.paragraph_format.space_after = Pt(2)
    base.paragraph_format.space_before = Pt(0)

    # Tighter style for the bullet list (otherwise it inherits Word's default
    # 8pt-after spacing which spreads bullets apart).
    try:
        bullet_style = doc.styles["List Bullet"]
        bullet_style.paragraph_format.space_after = Pt(1)
        bullet_style.paragraph_format.space_before = Pt(0)
    except KeyError:
        pass

    # Available width (page width minus margins): 8.5" - 0.6" - 0.6" = 7.3".
    # Use this as the right-tab stop position so dates align flush right.
    USABLE_WIDTH = Inches(7.3)

    def _add_bottom_border(paragraph) -> None:
        """Add a thin grey rule below the given paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")     # half-points -> 0.75pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Header block
    if rendered.name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(rendered.name)
        run.bold = True
        run.font.size = Pt(18)
    if rendered.headline:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(rendered.headline)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    contact_bits = [b for b in (rendered.email, rendered.location) if b]
    if contact_bits:
        p = doc.add_paragraph(" | ".join(contact_bits))
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _add_bottom_border(p)

    def _section_heading(title: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        _add_bottom_border(p)

    for section in rendered.section_order:
        if section == "summary" and rendered.summary:
            _section_heading("Summary")
            for para in rendered.summary.split("\n\n"):
                ps = doc.add_paragraph(para.strip())
                ps.paragraph_format.space_after = Pt(2)
        elif section == "experience" and rendered.experience:
            _section_heading("Experience")
            for it in rendered.experience:
                head = " | ".join(b for b in (it.title, it.company) if b)
                # Role line with a right-aligned tab stop for the dates.
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.tab_stops.add_tab_stop(
                    USABLE_WIDTH, WD_TAB_ALIGNMENT.RIGHT,
                )
                if head:
                    head_run = p.add_run(head)
                    head_run.bold = True
                if it.dates:
                    dates_run = p.add_run("\t" + it.dates)
                    dates_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    dates_run.font.size = Pt(9.5)
                if it.technologies:
                    tp = doc.add_paragraph()
                    tp.paragraph_format.space_after = Pt(1)
                    tech_run = tp.add_run(", ".join(it.technologies))
                    tech_run.italic = True
                    tech_run.font.size = Pt(9.5)
                    tech_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                for b in it.bullets:
                    doc.add_paragraph(b, style="List Bullet")
        elif section == "skills" and (rendered.skill_groups or rendered.skills):
            _section_heading("Skills")
            if rendered.skill_groups:
                for g in rendered.skill_groups:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    cat_run = p.add_run(f"{g.category}: ")
                    cat_run.bold = True
                    p.add_run(", ".join(g.skills))
            else:
                p = doc.add_paragraph(" · ".join(rendered.skills))
                p.paragraph_format.space_after = Pt(2)
        elif section == "certifications" and rendered.certifications:
            _section_heading("Certifications")
            for c in rendered.certifications:
                meta = ", ".join(b for b in (c.issuer, c.year) if b)
                line = f"{c.name}" + (f" ({meta})" if meta else "")
                doc.add_paragraph(line, style="List Bullet")
        elif section == "education" and rendered.education:
            _section_heading("Education")
            for ed in rendered.education:
                head_bits = [b for b in (ed.degree, ed.institution, ed.year) if b]
                head = ", ".join(head_bits)
                if ed.gpa:
                    head = f"{head} | GPA: {ed.gpa}"
                doc.add_paragraph(head, style="List Bullet")
                for h in ed.honors:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(f"· {h}")
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Discreet preview footer.
    if rendered.banner:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(rendered.banner)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF (reportlab) ──────────────────────────────────────────────────────────

def render_pdf(rendered: RenderedResume) -> bytes:
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, KeepTogether, ListFlowable, ListItem, Paragraph,
        SimpleDocTemplate, Table, TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        # Tighter margins (0.6" / 0.5") fit more content per page without
        # crowding the edges; common one-page-resume convention.
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        title=rendered.name or "Resume",
    )

    # All text uses the Helvetica family (matching the HTML stylesheet and the
    # DOCX Calibri choice). ReportLab's getSampleStyleSheet()["Normal"] defaults
    # to Times-Roman, so paragraph styles inherit fontName explicitly.
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "Name", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=18, leading=21,
        textColor="#1a1a1a", alignment=TA_LEFT, spaceAfter=1,
    )
    headline_style = ParagraphStyle(
        "Headline", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=10, leading=12,
        textColor="#555555", spaceAfter=1,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9.5, leading=11,
        textColor="#666666", spaceAfter=2,
    )
    dates_style = ParagraphStyle(
        "Dates", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9.5, leading=12,
        textColor="#666666", alignment=TA_RIGHT,
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=10.5, leading=12,
        textColor="#333333", spaceBefore=8, spaceAfter=2,
    )
    role_style = ParagraphStyle(
        "Role", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=10.5, leading=12,
        spaceBefore=5, spaceAfter=0,
    )
    tech_style = ParagraphStyle(
        "Tech", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=9.5, leading=11,
        textColor="#555555", spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=12.5,
        textColor="#222222",
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body_style, spaceAfter=1,
    )
    banner_style = ParagraphStyle(
        "Banner", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=8.5, leading=10,
        textColor="#888888", spaceBefore=10, spaceAfter=0,
    )

    # Page width minus margins (LETTER 8.5" - 0.6" - 0.6" = 7.3")
    USABLE_WIDTH = 7.3 * inch
    DATES_COL = 1.3 * inch
    TITLE_COL = USABLE_WIDTH - DATES_COL

    def _hr() -> HRFlowable:
        return HRFlowable(width="100%", thickness=0.5, color="#cccccc",
                          spaceBefore=0, spaceAfter=4)

    def _section_hr() -> HRFlowable:
        return HRFlowable(width="100%", thickness=0.5, color="#dddddd",
                          spaceBefore=0, spaceAfter=2)

    story: list = []

    def P(text: str, style):
        return Paragraph(_esc(text), style)

    if rendered.name:
        story.append(P(rendered.name, name_style))
    if rendered.headline:
        story.append(P(rendered.headline, headline_style))
    contact_bits = [b for b in (rendered.email, rendered.location) if b]
    if contact_bits:
        story.append(P(" &middot; ".join(_esc(b) for b in contact_bits), contact_style))
    if rendered.name or rendered.headline or contact_bits:
        story.append(_hr())

    section_h_label = {
        "summary": "Summary",
        "experience": "Experience",
        "skills": "Skills",
        "certifications": "Certifications",
        "education": "Education",
    }

    for section in rendered.section_order:
        if section == "summary" and rendered.summary:
            story.append(P(section_h_label[section].upper(), section_style))
            for para in rendered.summary.split("\n\n"):
                if para.strip():
                    story.append(P(para.strip(), body_style))
        elif section == "experience" and rendered.experience:
            story.append(P(section_h_label[section].upper(), section_style))
            story.append(_section_hr())
            for it in rendered.experience:
                head = " | ".join(b for b in (it.title, it.company) if b)
                block: list = []
                # Two-column role/dates row: title left, dates flush-right.
                if head or it.dates:
                    cells = [[
                        P(head, role_style) if head else "",
                        P(it.dates, dates_style) if it.dates else "",
                    ]]
                    t = Table(cells, colWidths=[TITLE_COL, DATES_COL])
                    t.setStyle(TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]))
                    block.append(t)
                if it.technologies:
                    block.append(P(", ".join(it.technologies), tech_style))
                if it.bullets:
                    block.append(ListFlowable(
                        [ListItem(P(b, bullet_style), leftIndent=10,
                                  value="bullet", bulletColor="#444444")
                         for b in it.bullets],
                        bulletType="bullet", start="•",
                        leftIndent=10, bulletFontSize=8,
                    ))
                if block:
                    story.append(KeepTogether(block))
        elif section == "skills" and (rendered.skill_groups or rendered.skills):
            story.append(P(section_h_label[section].upper(), section_style))
            story.append(_section_hr())
            if rendered.skill_groups:
                for g in rendered.skill_groups:
                    # Bold category prefix, then the skills.
                    line = (
                        f"<b>{_esc(g.category)}:</b> "
                        f"{' &middot; '.join(_esc(s) for s in g.skills)}"
                    )
                    story.append(Paragraph(line, body_style))
            else:
                story.append(P(" &middot; ".join(_esc(s) for s in rendered.skills),
                               body_style))
        elif section == "certifications" and rendered.certifications:
            story.append(P(section_h_label[section].upper(), section_style))
            story.append(_section_hr())
            items = []
            for c in rendered.certifications:
                meta = ", ".join(b for b in (c.issuer, c.year) if b)
                line = c.name + (f" ({meta})" if meta else "")
                items.append(ListItem(P(line, bullet_style),
                                      leftIndent=10, value="bullet"))
            story.append(ListFlowable(items, bulletType="bullet",
                                      start="•", leftIndent=10,
                                      bulletFontSize=8))
        elif section == "education" and rendered.education:
            story.append(P(section_h_label[section].upper(), section_style))
            story.append(_section_hr())
            items = []
            for ed in rendered.education:
                head_bits = [b for b in (ed.degree, ed.institution, ed.year) if b]
                head = ", ".join(head_bits)
                if ed.gpa:
                    head = f"{head}  |  GPA: {ed.gpa}"
                # If honors are present, put them under the entry as a sub-list.
                if ed.honors:
                    sub = ListFlowable(
                        [ListItem(P(h, bullet_style), leftIndent=10,
                                  value="bullet", bulletColor="#888888")
                         for h in ed.honors],
                        bulletType="bullet", start="·",
                        leftIndent=20, bulletFontSize=7,
                    )
                    items.append(ListItem(
                        KeepTogether([P(head, bullet_style), sub]),
                        leftIndent=10, value="bullet",
                    ))
                else:
                    items.append(ListItem(P(head, bullet_style),
                                          leftIndent=10, value="bullet"))
            story.append(ListFlowable(items, bulletType="bullet",
                                      start="•", leftIndent=10,
                                      bulletFontSize=8))

    # Discreet preview footer.
    if rendered.banner:
        story.append(Paragraph(_esc(rendered.banner), banner_style))

    doc.build(story)
    return buf.getvalue()


def _esc(s: str) -> str:
    """ReportLab Paragraph uses minimal HTML; escape the same five entities."""
    return (str(s)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


# ── Public dispatch ─────────────────────────────────────────────────────────


ExportFormat = Literal["md", "txt", "html", "json", "docx", "pdf"]


_CONTENT_TYPE = {
    "md":   "text/markdown; charset=utf-8",
    "txt":  "text/plain; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
}


def export_content_type(fmt: str) -> str:
    return _CONTENT_TYPE.get(fmt, "application/octet-stream")


def export_file_extension(fmt: str) -> str:
    return "json" if fmt == "json" else fmt


def render(format: str, rendered: RenderedResume) -> bytes:
    """Single entry point for the API: render the resume in the requested format
    and return raw bytes (text formats are utf-8 encoded)."""
    if format == "md":
        return render_markdown(rendered).encode("utf-8")
    if format == "txt":
        return render_plain_text(rendered).encode("utf-8")
    if format == "html":
        return render_html(rendered).encode("utf-8")
    if format == "json":
        return json.dumps(render_json_resume(rendered), indent=2).encode("utf-8")
    if format == "docx":
        return render_docx(rendered)
    if format == "pdf":
        return render_pdf(rendered)
    raise ValueError(f"Unknown export format: {format!r}")
